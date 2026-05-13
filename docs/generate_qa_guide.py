from docx import Document
from docx.shared import Pt
from pathlib import Path

doc = Document()
doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(11)

doc.add_heading('SQALogic Test Site — QA Automation User Guide', level=0)
doc.add_paragraph('Audience: QA automation engineers  |  Version: 0.1.0  |  Date: 2026-04-14')

doc.add_heading('1. Purpose', level=1)
doc.add_paragraph(
    'This guide explains how to automate tests against the SQALogic Test Site. The site '
    'is intentionally designed to stress-test automation frameworks: element IDs, CSS '
    'classes, DOM order, rendered roles, and timing all mutate across "releases". A '
    'locator that works today may fail tomorrow. Your tests must be resilient.'
)

doc.add_heading('2. Environment', level=1)
tbl = doc.add_table(rows=1, cols=2); tbl.style = 'Light Grid Accent 1'
tbl.rows[0].cells[0].text = 'Item'; tbl.rows[0].cells[1].text = 'Value'
for k, v in [
    ('Base URL (local)', 'http://localhost:3000'),
    ('Base URL (preview)', 'Vercel preview URL assigned per deployment'),
    ('Backend', 'Supabase (shared state — release counter is global)'),
    ('Demo account', 'demo@sqalogic.ca / demo123'),
    ('Recommended browser', 'Chromium (Playwright / Vibium)'),
]:
    r = tbl.add_row().cells; r[0].text = k; r[1].text = v

doc.add_heading('3. Release Mutation Model (critical)', level=1)
doc.add_paragraph(
    'The site carries a global integer "release" stored in the release_state table in '
    'Supabase. Incrementing it changes the DOM in the following ways:'
)
for bullet in [
    'Dynamic IDs: inputs receive IDs like "login_email_v{release}_{hash}" — they change every release.',
    'Dynamic classes: utility classes append a per-release hash suffix.',
    'Attribute flip: release ≥ 3 removes id and data-testid; release ≥ 4 adds data-qa="...".',
    'Role swap (login): release ≥ 2 (even) swaps the DOM order of email and password fields.',
    'Fake button: release ≥ 3 renders the Search submit as <div role="button"> instead of <button>.',
    'CTA label: release toggles between "Find Flights" and "Search flights".',
    'Random delays: release ≥ 2 introduces random async delays up to ~150 + 100 × release ms (capped at 800 ms).',
]:
    doc.add_paragraph(bullet, style='List Bullet')
doc.add_paragraph(
    'Automation implication: NEVER rely on raw id, class, or DOM index. Locate by stable '
    'semantics (label text, visible text, role+name) or by anchoring from a nearby label.'
)

doc.add_heading('4. Recommended Locator Strategy', level=1)
for bullet in [
    'Prefer getByLabel / getByRole with accessible name — they survive id and class mutation.',
    'For the Search submit: match by visible text regex /Find Flights|Search flights/ and do NOT require tagName=button (could be div[role="button"]).',
    'For login: always read the field label ("Email address" / "Password"), never positional order — fields swap.',
    'Avoid CSS selectors that depend on generated hashes (e.g. .btn-x9k2) — they rotate each release.',
    'If an element has data-qa, use it; otherwise fall back to role + text.',
    'Use retry/wait helpers tolerant to ~800 ms delays.',
]:
    doc.add_paragraph(bullet, style='List Bullet')

doc.add_heading('5. Core User Journeys to Automate', level=1)
journeys = [
    ('Registration', '/register', 'Sign up with name + email + password (≥6 chars); duplicates rejected.'),
    ('Login',        '/login',    'Sign in with email + password; redirects to /search. Demo credentials available.'),
    ('Search',       '/search',   'Select trip type, from, to, date, passengers; submit → /results?query.'),
    ('Results',      '/results',  'List of flights parsed from query string; choose a flight → /book.'),
    ('Book',         '/book',     'Enter passenger, choose seat (SeatMap), select baggage → /payment.'),
    ('Payment',      '/payment',  'Submit payment details → /confirmation.'),
    ('Confirmation', '/confirmation', 'Booking recorded in Supabase; verify presence in /my-trips.'),
    ('My Trips',     '/my-trips', 'Authenticated list of past bookings for the logged-in user.'),
    ('Logout',       'Header',    'Clears session from localStorage (sqa_user).'),
    ('Embed mode',   '/embed',    'Widget embedded in iframe — test cross-frame locators.'),
]
t = doc.add_table(rows=1, cols=3); t.style = 'Light Grid Accent 1'
h = t.rows[0].cells; h[0].text = 'Journey'; h[1].text = 'Route'; h[2].text = 'Notes'
for j, r_, n in journeys:
    row = t.add_row().cells; row[0].text = j; row[1].text = r_; row[2].text = n

doc.add_heading('6. Test Data & State Reset', level=1)
for bullet in [
    'Users are stored in Supabase table users; passwords are SHA-256 hashed.',
    'Bookings persist in the bookings table keyed by user_email.',
    'Session is stored client-side in localStorage under "sqa_user" — clear it between tests for isolation.',
    'release_state has a single row (id=1). Bumping it affects ALL concurrent sessions — coordinate shared test environments.',
    'Use unique emails per test (e.g. qa+{timestamp}@sqalogic.ca) to avoid "Email already registered" collisions.',
]:
    doc.add_paragraph(bullet, style='List Bullet')

doc.add_heading('7. Handling Flakiness', level=1)
for bullet in [
    'Randomized delays — always use web-first assertions / auto-wait; do NOT hardcode sleeps below 1 s.',
    'Real-time release updates — the app subscribes to postgres_changes; the DOM may mutate mid-test. Re-query locators after navigation.',
    'Seat map — rendered asynchronously; wait for at least one seat to be visible before clicking.',
    'Network — Supabase calls may be slow on first load. Expect a cold-start on the first test in a run.',
]:
    doc.add_paragraph(bullet, style='List Bullet')

doc.add_heading('8. Sample Playwright Snippet', level=1)
p = doc.add_paragraph()
run = p.add_run(
"""import { test, expect } from '@playwright/test';

test('search → book happy path', async ({ page }) => {
  await page.goto('http://localhost:3000/login');
  await page.getByLabel('Email address').fill('demo@sqalogic.ca');
  await page.getByLabel('Password').fill('demo123');
  await page.getByRole('button', { name: /sign in/i }).click();

  await expect(page).toHaveURL(/\\/search/);
  await page.getByLabel('From').selectOption({ label: 'YUL - Montreal' });
  await page.getByLabel('To').selectOption({ label: 'JFK - New York' });

  // CTA text and tag can mutate — match by accessible name only.
  await page.getByRole('button', { name: /find flights|search flights/i }).click();

  await expect(page).toHaveURL(/\\/results/);
  await page.getByRole('link', { name: /select/i }).first().click();

  await page.getByLabel(/passenger/i).fill('QA Bot');
  await page.locator('[data-seat]').first().click();
  await page.getByRole('button', { name: /continue|next|pay/i }).click();
});
"""
)
run.font.name = 'Consolas'; run.font.size = Pt(9)

doc.add_heading('9. Sample Vibium Snippet', level=1)
p = doc.add_paragraph()
run = p.add_run(
"""await browser_start({ headless: true });
await browser_navigate({ url: 'http://localhost:3000/login' });
await browser_fill({ locator: { label: 'Email address' }, value: 'demo@sqalogic.ca' });
await browser_fill({ locator: { label: 'Password' },      value: 'demo123' });
await browser_click({ locator: { role: 'button', name: /sign in/i } });
await browser_wait_for_url({ url: /\\/search/ });
await browser_click({ locator: { role: 'button', name: /find flights|search flights/i } });
"""
)
run.font.name = 'Consolas'; run.font.size = Pt(9)

doc.add_heading('10. Reporting Issues', level=1)
for bullet in [
    'Capture screenshots on failure (browser_screenshot / page.screenshot).',
    'Record the current release value (read from release_state or window state) in the test report — bug reproduction often depends on it.',
    'Include user email used, timestamp, and full URL including query string.',
]:
    doc.add_paragraph(bullet, style='List Bullet')

doc.add_heading('11. Do Not', level=1)
for bullet in [
    'Do not rely on generated id/class values — they rotate.',
    'Do not bump the release during a shared test run — it affects every user.',
    'Do not commit real credentials; the demo account is the only shared identity.',
    'Do not assume the Search CTA is a <button> element — use role queries.',
]:
    doc.add_paragraph(bullet, style='List Bullet')

output = Path(r'C:/Users/Luis/.gemini/antigravity/scratch/SQATest/sqalogic-test-site/docs/SQALogic-QA-Automation-Guide.docx')
doc.save(output)
print(f'Saved: {output}')
