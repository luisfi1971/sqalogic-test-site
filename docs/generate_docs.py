from docx import Document
from docx.shared import Pt, RGBColor
from pathlib import Path

doc = Document()

styles = doc.styles
normal = styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)

title = doc.add_heading('SQALogic Test Site — Project Documentation', level=0)

doc.add_paragraph('Version: 0.1.0  |  Framework: Next.js 16.2.3  |  Date: 2026-04-14')

doc.add_heading('1. Overview', level=1)
doc.add_paragraph(
    'The SQALogic Test Site is a Next.js web application built to demonstrate '
    'a flight booking flow. It uses the Next.js App Router, React 19, Tailwind CSS 4, '
    'and Supabase as the backend for data persistence and authentication. The project '
    'is deployed on Vercel.'
)
doc.add_paragraph(
    'IMPORTANT: This project uses Next.js 16, which introduces breaking changes compared '
    'to earlier versions. APIs, conventions, and file structure may differ from typical '
    'Next.js training data. Always consult node_modules/next/dist/docs/ before modifying code.'
)

doc.add_heading('2. Tech Stack', level=1)
stack = [
    ('Framework', 'Next.js 16.2.3 (App Router)'),
    ('UI Library', 'React 19.2.4 + React DOM 19.2.4'),
    ('Styling', 'Tailwind CSS 4 with @tailwindcss/postcss'),
    ('Backend', 'Supabase (@supabase/supabase-js ^2.103.0)'),
    ('Language', 'TypeScript 5'),
    ('Linting', 'ESLint 9 with eslint-config-next'),
    ('Deployment', 'Vercel'),
]
table = doc.add_table(rows=1, cols=2)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = 'Layer'
hdr[1].text = 'Technology'
for layer, tech in stack:
    row = table.add_row().cells
    row[0].text = layer
    row[1].text = tech

doc.add_heading('3. Project Structure', level=1)
structure = """sqalogic-test-site/
├── app/                    # Next.js App Router pages and routes
│   ├── book/               # Booking flow page
│   ├── components/         # Shared React components
│   │   ├── Calendar.tsx
│   │   ├── Chrome.tsx
│   │   ├── ConfirmModal.tsx
│   │   ├── Dropdown.tsx
│   │   ├── Header.tsx
│   │   ├── SeatMap.tsx
│   │   ├── ShadowInput.tsx
│   │   └── Tooltip.tsx
│   ├── confirmation/       # Booking confirmation page
│   ├── embed/              # Embeddable widget route
│   ├── lib/                # Shared utilities
│   │   ├── hash.ts         # Hashing helpers
│   │   └── supabase.ts     # Supabase client
│   ├── login/              # User login page
│   ├── my-trips/           # User trip history
│   ├── payment/            # Payment step
│   ├── register/           # User registration
│   ├── results/            # Flight search results
│   ├── search/             # Search form
│   ├── globals.css         # Global Tailwind styles
│   ├── layout.tsx          # Root layout
│   ├── page.tsx            # Landing page
│   └── providers.tsx       # Context providers
├── public/                 # Static assets
├── scripts/
│   └── migrate.mjs         # DB migration runner
├── supabase/
│   └── migrations/
│       ├── 0001_initial.sql
│       └── 0002_seat_baggage_release.sql
├── .env.local              # Local environment variables
├── next.config.ts
├── package.json
├── tsconfig.json
└── README.md
"""
p = doc.add_paragraph()
run = p.add_run(structure)
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_heading('4. Application Flow', level=1)
doc.add_paragraph(
    'The site implements a complete flight booking journey composed of the following pages:'
)
flow = [
    ('/', 'Landing page with entry point to search.'),
    ('/search', 'Flight search form — origin, destination, dates, passengers.'),
    ('/results', 'Search results list with filtering and selection.'),
    ('/book', 'Booking details — passenger info, seat selection (SeatMap), baggage.'),
    ('/payment', 'Payment capture step.'),
    ('/confirmation', 'Post-booking confirmation screen.'),
    ('/login', 'User authentication.'),
    ('/register', 'New user signup.'),
    ('/my-trips', 'Authenticated user trip history.'),
    ('/embed', 'Embeddable version of the widget for third-party sites.'),
]
t2 = doc.add_table(rows=1, cols=2)
t2.style = 'Light Grid Accent 1'
h2 = t2.rows[0].cells
h2[0].text = 'Route'
h2[1].text = 'Purpose'
for route, purpose in flow:
    r = t2.add_row().cells
    r[0].text = route
    r[1].text = purpose

doc.add_heading('5. Database', level=1)
doc.add_paragraph(
    'The backend is Supabase (PostgreSQL). Schema is managed through SQL migration files '
    'stored in supabase/migrations/ and executed by scripts/migrate.mjs.'
)
doc.add_paragraph('Current migrations:', style='Intense Quote')
doc.add_paragraph('• 0001_initial.sql — base schema (users, flights, bookings).', style='List Bullet')
doc.add_paragraph('• 0002_seat_baggage_release.sql — adds seat selection and baggage support.', style='List Bullet')

doc.add_heading('6. Environment Setup', level=1)
doc.add_paragraph('Required environment variables (.env.local):')
env_tbl = doc.add_table(rows=1, cols=2)
env_tbl.style = 'Light Grid Accent 1'
eh = env_tbl.rows[0].cells
eh[0].text = 'Variable'
eh[1].text = 'Purpose'
for k, v in [
    ('NEXT_PUBLIC_SUPABASE_URL', 'Supabase project URL (client-side).'),
    ('NEXT_PUBLIC_SUPABASE_ANON_KEY', 'Supabase anon public key (client-side).'),
    ('SUPABASE_SERVICE_ROLE_KEY', 'Service role key for migrations (server only).'),
]:
    r = env_tbl.add_row().cells
    r[0].text = k
    r[1].text = v

doc.add_heading('7. Scripts', level=1)
for name, desc in [
    ('npm run dev', 'Start the Next.js development server at http://localhost:3000.'),
    ('npm run build', 'Produce an optimized production build.'),
    ('npm run start', 'Run the production server after build.'),
    ('npm run lint', 'Run ESLint across the codebase.'),
    ('npm run migrate', 'Apply pending Supabase migrations via scripts/migrate.mjs.'),
]:
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(name)
    r.bold = True
    p.add_run(f' — {desc}')

doc.add_heading('8. Deployment', level=1)
doc.add_paragraph(
    'The project is linked to Vercel (see .vercel/). Pushes to the main branch trigger '
    'preview/production deployments. Ensure all environment variables are configured in '
    'the Vercel project settings before deploying.'
)

doc.add_heading('9. Development Notes', level=1)
doc.add_paragraph(
    'Because Next.js 16 introduces breaking changes, contributors must read the relevant '
    'guide in node_modules/next/dist/docs/ before writing code — especially for routing, '
    'caching, Server Components, and Server Actions. Heed all deprecation notices emitted '
    'during build.'
)

output_dir = Path(r'C:/Users/Luis/.gemini/antigravity/scratch/SQATest/sqalogic-test-site/docs')
output_dir.mkdir(parents=True, exist_ok=True)
output = output_dir / 'SQALogic-Test-Site-Documentation.docx'
doc.save(output)
print(f'Saved: {output}')
